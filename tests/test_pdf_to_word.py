"""Tests for the pdf_to_docx service (PDF -> editable .docx)."""

import io
import zipfile
from pathlib import Path

import pymupdf
import pytest

from app.api_errors import ApiError
from app.services.pdf_tools import _docx_is_effectively_empty, pdf_to_docx


def _text_pdf(pages: int = 1, text: str = "Contrato de teste 12345") -> bytes:
    doc = pymupdf.open()
    for _ in range(pages):
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=14)
    out = doc.tobytes()
    doc.close()
    return out


def _blank_pdf(pages: int = 1) -> bytes:
    doc = pymupdf.open()
    for _ in range(pages):
        doc.new_page()  # no text -> simulates a scanned/image-only page
    out = doc.tobytes()
    doc.close()
    return out


def _encrypted_pdf() -> bytes:
    doc = pymupdf.open()
    doc.new_page().insert_text((72, 72), "segredo", fontsize=12)
    out = doc.tobytes(encryption=pymupdf.PDF_ENCRYPT_AES_256, user_pw="pw", owner_pw="pw")
    doc.close()
    return out


def test_text_pdf_returns_valid_docx():
    result = pdf_to_docx(_text_pdf(text="Contrato de teste 12345"))
    assert isinstance(result, bytes) and len(result) > 0
    with zipfile.ZipFile(io.BytesIO(result)) as zf:
        names = zf.namelist()
        assert "word/document.xml" in names
        document_xml = zf.read("word/document.xml").decode("utf-8", "replace")
    assert "12345" in document_xml


def test_scanned_pdf_raises_422():
    with pytest.raises(ApiError) as exc:
        pdf_to_docx(_blank_pdf())
    assert exc.value.status_code == 422
    assert exc.value.code == "scanned_pdf"


def test_encrypted_pdf_raises_400():
    with pytest.raises(ApiError) as exc:
        pdf_to_docx(_encrypted_pdf())
    assert exc.value.status_code == 400
    assert exc.value.code == "password_protected_pdf"


def test_corrupt_bytes_raise_400():
    with pytest.raises(ApiError) as exc:
        pdf_to_docx(b"this is not a pdf")
    assert exc.value.status_code == 400
    assert exc.value.code == "invalid_pdf"


def test_empty_docx_helper(tmp_path: Path):
    import docx

    blank = tmp_path / "blank.docx"
    docx.Document().save(blank)
    assert _docx_is_effectively_empty(blank) is True

    filled = tmp_path / "filled.docx"
    d = docx.Document()
    d.add_paragraph("Olá mundo")
    d.save(filled)
    assert _docx_is_effectively_empty(filled) is False


def test_v2_pdf_to_word_text_returns_200(client):
    pdf_bytes = io.BytesIO(_text_pdf(text="Relatorio 98765"))
    response = client.post(
        "/v2/pdf-to-word",
        files={"file": ("test.pdf", pdf_bytes, "application/pdf")},
        data={"options": "{}"},
    )
    assert response.status_code == 200
    assert response.headers["content-type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.headers["content-disposition"].endswith('.docx"')
    with zipfile.ZipFile(io.BytesIO(response.content)) as zf:
        assert "word/document.xml" in zf.namelist()


def test_v2_pdf_to_word_scanned_returns_422_envelope(client):
    response = client.post(
        "/v2/pdf-to-word",
        files={"file": ("scan.pdf", io.BytesIO(_blank_pdf()), "application/pdf")},
        data={"options": "{}"},
    )
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "scanned_pdf"


def test_v2_pdf_to_word_encrypted_returns_400_envelope(client):
    response = client.post(
        "/v2/pdf-to-word",
        files={"file": ("enc.pdf", io.BytesIO(_encrypted_pdf()), "application/pdf")},
        data={"options": "{}"},
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "password_protected_pdf"
