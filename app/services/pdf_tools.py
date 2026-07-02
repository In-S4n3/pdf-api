"""Shared PDF processing services for v1 and v2 routes."""

from __future__ import annotations

import hashlib
import io
import logging
import re as _re
import subprocess
import tempfile
import threading
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import img2pdf
import regex

from app.api_errors import ApiError

logger = logging.getLogger(__name__)

MAX_PAGES = 200  # coarse guard; the 55s subprocess timeout is the real runtime bound

# pdf_to_xlsx caps (in-process; see 2026-07-02-pdf-para-excel-design.md §3.2).
MAX_TABLES = 200  # cap total worksheets
MAX_CELLS = 500_000  # cap total cells written — bounds the in-memory openpyxl workbook
MAX_PATHS_PER_PAGE = 5000  # reject vector-graphics-bomb pages before find_tables clustering

OFFICE_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
}

IMAGE_MIMES = {
    "image/jpeg",
    "image/png",
    "image/tiff",
}

MIME_TO_EXT = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": ".pptx",
}

IMAGE_MIME_TO_EXT = {
    "image/jpeg": ".jpg",
    "image/png": ".png",
    "image/tiff": ".tiff",
}

EXTENSION_TO_MIME = {
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
}

LANGUAGE_MAP = {
    "english": "eng",
    "spanish": "spa",
    "french": "fra",
    "german": "deu",
    "portuguese": "por",
    "italian": "ita",
    "chinese": "chi_sim",
    "jpn": "jpn",
}

CONFORMANCE_MAP = {
    "pdfa-1b": "1",
    "pdfa-2b": "2",
    "pdfa-3b": "3",
}

REGEX_TIMEOUT_SECONDS = 0.5

# Updated patterns — \b anchors prevent partial matches like "fxxx@yyy" capture
EMAIL_PATTERN = r"\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b"
PHONE_PATTERN = r"\b\+?\(?\d{1,4}\)?[\d\s./\-]{6,14}\d\b"
PATTERNS = {
    "email": EMAIL_PATTERN,
    "phone": PHONE_PATTERN,
}

VALID_REDACTION_STRATEGIES = ("email", "phone", "custom", "regex")
MAX_REGEX_LENGTH = 500

REDACTION_LOCK = threading.Lock()


@dataclass(frozen=True, slots=True)
class RedactionMatch:
    id: str
    page: int            # 0-indexed
    bbox: tuple[float, float, float, float]  # (x0, y0, x1, y1) in PDF points
    kind: str            # 'email' | 'phone' | 'custom' | 'regex'
    context: str         # word text inside the bbox (visible in UI)
    full_match: str      # the entire regex match (may span multiple words)


def _make_match(
    strategy: str,
    page_idx: int,
    bbox: tuple[float, float, float, float],
    context: str,
    full_match: str,
) -> RedactionMatch:
    """Build a RedactionMatch with a deterministic 16-char hex ID.

    Same (strategy, page, bbox, context) -> same ID, so a confirmed-IDs
    round-trip from frontend to backend continues to identify the same
    matches even though the helper is called twice (once for preview,
    once for apply).
    """
    digest = hashlib.sha1(
        f"{strategy}|{page_idx}|{bbox[0]:.2f},{bbox[1]:.2f},{bbox[2]:.2f},{bbox[3]:.2f}|{context}".encode()
    ).hexdigest()[:16]
    return RedactionMatch(
        id=digest, page=page_idx, bbox=bbox, kind=strategy,
        context=context, full_match=full_match,
    )


def _compile_pattern(
    strategy: str, custom_text: str, regex_pattern: str
) -> tuple[str, int]:
    """Validate inputs and return (pattern, flags). Raises ApiError on bad input."""
    if strategy not in VALID_REDACTION_STRATEGIES:
        raise ApiError(400, "invalid_redaction_strategy", "Estratégia de redacção inválida.")

    if strategy == "custom":
        if not custom_text.strip():
            raise ApiError(400, "missing_custom_text", "Texto personalizado é obrigatório.")
        return _re.escape(custom_text), regex.IGNORECASE

    if strategy == "regex":
        if not regex_pattern.strip():
            raise ApiError(400, "missing_regex_pattern", "Padrão regex é obrigatório.")
        if len(regex_pattern) > MAX_REGEX_LENGTH:
            raise ApiError(
                400, "regex_too_long",
                f"Padrão regex demasiado longo (máx {MAX_REGEX_LENGTH} caracteres).",
            )
        try:
            regex.compile(regex_pattern)
        except regex.error as exc:
            raise ApiError(400, "invalid_regex_pattern", "Padrão regex inválido.") from exc
        return regex_pattern, 0

    return PATTERNS[strategy], 0


def _extract_matches(
    doc, *, strategy: str, custom_text: str, regex_pattern: str
) -> list[RedactionMatch]:
    """Find all matches across all pages. Used by both /preview and /redact."""
    if doc.needs_pass:
        raise ApiError(
            400, "password_protected_pdf",
            "Este PDF está protegido por palavra-passe. Remova a proteção antes de redactar.",
        )
    pattern_str, flags = _compile_pattern(strategy, custom_text, regex_pattern)
    matches: list[RedactionMatch] = []

    for page_idx, page in enumerate(doc):
        text = page.get_text("text")
        try:
            iter_matches = list(
                regex.finditer(pattern_str, text, flags=flags, timeout=REGEX_TIMEOUT_SECONDS)
            )
        except regex.error as exc:
            raise ApiError(400, "invalid_regex_pattern", "Padrão regex inválido.") from exc
        except TimeoutError as exc:
            raise ApiError(
                400, "regex_too_slow",
                "O padrão regex é demasiado complexo (possível ReDoS). Simplifique-o.",
            ) from exc

        full_match_strings = {m.group() for m in iter_matches}

        for match_str in full_match_strings:
            for rect in page.search_for(match_str):
                # Decompose the match rect into per-word bboxes for clean preview highlights.
                words = page.get_text("words", clip=rect)
                if not words:
                    bbox = (rect.x0, rect.y0, rect.x1, rect.y1)
                    matches.append(_make_match(strategy, page_idx, bbox, match_str, match_str))
                    continue
                for w in words:
                    bbox = (w[0], w[1], w[2], w[3])
                    word_text = w[4]
                    matches.append(_make_match(strategy, page_idx, bbox, word_text, match_str))

    return matches


def _trim_process_output(value: str, limit: int = 500) -> str:
    value = value.strip()
    if len(value) <= limit:
        return value
    return value[-limit:]


def _run_command(
    command: list[str],
    *,
    timeout: int,
    missing_message: str,
) -> subprocess.CompletedProcess[str]:
    try:
        return subprocess.run(
            command,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout,
        )
    except FileNotFoundError as exc:
        raise ApiError(
            status_code=503,
            code="tool_unavailable",
            message=missing_message,
        ) from exc
    except subprocess.TimeoutExpired as exc:
        raise ApiError(
            status_code=504,
            code="processing_timeout",
            message="O processamento excedeu o tempo limite. Tente com um ficheiro mais pequeno.",
        ) from exc


def _resolve_convert_content_type(content_type: str | None, filename: str | None) -> str | None:
    if content_type:
        normalized = content_type.lower()
        if normalized in OFFICE_MIMES or normalized in IMAGE_MIMES:
            return normalized

    suffix = Path(filename or "").suffix.lower()
    return EXTENSION_TO_MIME.get(suffix)


def compress_pdf(content: bytes) -> bytes:
    """Compress a PDF using PyMuPDF."""
    import pymupdf

    doc = None
    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
        if doc.needs_pass:
            raise ApiError(
                status_code=400,
                code="password_protected_pdf",
                message="Este PDF está protegido por palavra-passe. Remova a proteção antes de comprimir.",
            )
        doc.rewrite_images(
            dpi_threshold=150,
            dpi_target=96,
            quality=75,
        )
        return doc.tobytes(
            garbage=4,
            deflate=True,
            clean=True,
            use_objstms=True,
        )
    except ApiError:
        raise
    except Exception as exc:
        raise ApiError(
            status_code=400,
            code="invalid_pdf",
            message="Nao foi possivel abrir o PDF. Verifique se o ficheiro e valido.",
        ) from exc
    finally:
        if doc is not None:
            doc.close()


def flatten_pdf(content: bytes) -> bytes:
    """Flatten annotations and form fields into the page content."""
    import pymupdf

    doc = None
    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
        if doc.needs_pass:
            raise ApiError(
                status_code=400,
                code="password_protected_pdf",
                message="Este PDF está protegido por palavra-passe. Remova a proteção antes de achatar.",
            )
        doc.bake(annots=True, widgets=True)
        return doc.tobytes(garbage=4, deflate=True)
    except ApiError:
        raise
    except Exception as exc:
        raise ApiError(
            status_code=400,
            code="invalid_pdf",
            message="Nao foi possivel abrir o PDF. Verifique se o ficheiro e valido.",
        ) from exc
    finally:
        if doc is not None:
            doc.close()


def _convert_office(content: bytes, content_type: str) -> bytes:
    ext = MIME_TO_EXT[content_type]

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / f"input{ext}"
        output_path = Path(tmpdir) / "input.pdf"
        input_path.write_bytes(content)

        result = _run_command(
            [
                "soffice",
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                tmpdir,
                f"-env:UserInstallation=file://{tmpdir}/profile",
                str(input_path),
            ],
            timeout=55,
            missing_message="LibreOffice não está disponível neste ambiente.",
        )

        if result.returncode != 0 or not output_path.exists():
            raise ApiError(
                status_code=500,
                code="conversion_failed",
                message="Falha na conversao do documento.",
                details={"stderr": _trim_process_output(result.stderr)},
            )

        return output_path.read_bytes()


def _convert_image_with_libreoffice(content: bytes, content_type: str) -> bytes:
    ext = IMAGE_MIME_TO_EXT.get(content_type, ".jpg")

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / f"input{ext}"
        output_path = Path(tmpdir) / "input.pdf"
        input_path.write_bytes(content)

        result = _run_command(
            [
                "soffice",
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                tmpdir,
                f"-env:UserInstallation=file://{tmpdir}/profile",
                str(input_path),
            ],
            timeout=55,
            missing_message="LibreOffice não está disponível neste ambiente.",
        )

        if result.returncode != 0 or not output_path.exists():
            raise ApiError(
                status_code=500,
                code="conversion_failed",
                message="Falha na conversao da imagem.",
                details={"stderr": _trim_process_output(result.stderr)},
            )

        return output_path.read_bytes()


def _sanitize_image(content: bytes) -> bytes:
    """Re-encode an image to strip ICC profiles and problematic metadata.

    Some ICC profiles (e.g. short lcms "c2" profiles) cause Adobe Acrobat to
    render the resulting PDF as a solid black page.  Re-saving through Pillow
    in sRGB without the original ICC profile eliminates the issue.
    """
    from PIL import Image

    img = Image.open(io.BytesIO(content))
    if not img.info.get("icc_profile"):
        return content  # nothing to strip

    out = io.BytesIO()
    save_kwargs: dict[str, object] = {}
    if img.format == "JPEG" or img.mode in ("RGB", "L"):
        save_kwargs["format"] = "JPEG"
        save_kwargs["quality"] = 95
        if img.mode == "RGBA":
            img = img.convert("RGB")
    else:
        save_kwargs["format"] = img.format or "PNG"
        if img.mode == "RGBA":
            background = Image.new("RGB", img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background

    img.save(out, **save_kwargs)
    return out.getvalue()


def convert_to_pdf(content: bytes, content_type: str | None, filename: str | None) -> bytes:
    """Convert a supported office document or image to PDF."""
    resolved_content_type = _resolve_convert_content_type(content_type, filename)
    if resolved_content_type is None:
        raise ApiError(
            status_code=400,
            code="unsupported_media_type",
            message="Formato nao suportado.",
        )

    if resolved_content_type in OFFICE_MIMES:
        return _convert_office(content, resolved_content_type)

    try:
        return img2pdf.convert(_sanitize_image(content))
    except Exception:
        return _convert_image_with_libreoffice(content, resolved_content_type)


def pdf_first_page_to_image(content: bytes, fmt: str) -> tuple[bytes, str, str]:
    """Render the first page of a PDF to PNG or JPEG."""
    import pymupdf

    doc = None
    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ApiError(
            status_code=400,
            code="invalid_pdf",
            message="Nao foi possivel abrir o PDF. Verifique se o ficheiro e valido.",
        ) from exc

    try:
        if len(doc) == 0:
            raise ApiError(
                status_code=400,
                code="invalid_pdf",
                message="O PDF nao contem paginas para converter.",
            )

        page = doc[0]
        pix = page.get_pixmap(dpi=300)
        if fmt == "jpeg":
            return pix.tobytes("jpeg", jpg_quality=92), "image/jpeg", "jpg"
        return pix.tobytes("png"), "image/png", "png"
    except ApiError:
        raise
    except Exception as exc:
        raise ApiError(
            status_code=500,
            code="conversion_failed",
            message="Nao foi possivel converter este PDF para imagem.",
        ) from exc
    finally:
        if doc is not None:
            doc.close()


MAX_PAGES_FOR_IMAGES = 20


def pdf_to_images(content: bytes, fmt: str) -> tuple[bytes, str, str]:
    """Render all pages of a PDF to images and return a ZIP archive."""
    import io
    import zipfile

    import pymupdf

    doc = None
    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ApiError(
            status_code=400,
            code="invalid_pdf",
            message="Nao foi possivel abrir o PDF. Verifique se o ficheiro e valido.",
        ) from exc

    try:
        page_count = len(doc)
        if page_count == 0:
            raise ApiError(
                status_code=400,
                code="invalid_pdf",
                message="O PDF nao contem paginas para converter.",
            )

        if page_count > MAX_PAGES_FOR_IMAGES:
            raise ApiError(
                status_code=400,
                code="too_many_pages",
                message=(
                    f"O PDF tem {page_count} paginas (maximo: {MAX_PAGES_FOR_IMAGES}). "
                    "Use a ferramenta Extrair Paginas para selecionar as paginas pretendidas."
                ),
            )

        ext = "jpg" if fmt == "jpeg" else "png"
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=300)
                if fmt == "jpeg":
                    img_bytes = pix.tobytes("jpeg", jpg_quality=92)
                else:
                    img_bytes = pix.tobytes("png")
                zf.writestr(f"pagina-{i + 1}.{ext}", img_bytes)

        return buf.getvalue(), "application/zip", "zip"
    except ApiError:
        raise
    except Exception as exc:
        raise ApiError(
            status_code=500,
            code="conversion_failed",
            message="Nao foi possivel converter este PDF para imagens.",
        ) from exc
    finally:
        if doc is not None:
            doc.close()


def ocr_pdf(content: bytes, language: str) -> bytes:
    """Run OCRmyPDF with the requested language."""
    lang_code = LANGUAGE_MAP.get(language)
    if lang_code is None:
        raise ApiError(
            status_code=400,
            code="unsupported_language",
            message=(
                f"Idioma não suportado: {language}. "
                f"Suportados: {', '.join(LANGUAGE_MAP.keys())}"
            ),
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.pdf"
        output_path = Path(tmpdir) / "output.pdf"
        input_path.write_bytes(content)

        result = _run_command(
            [
                "ocrmypdf",
                "--skip-text",
                "--output-type",
                "pdf",
                "-l",
                lang_code,
                "--deskew",
                "--clean",
                "--optimize",
                "1",
                "--tesseract-timeout",
                "60",
                str(input_path),
                str(output_path),
            ],
            timeout=55,
            missing_message="OCRmyPDF não está disponível neste ambiente.",
        )

        if result.returncode == 2:
            raise ApiError(
                status_code=400,
                code="invalid_pdf",
                message="Ficheiro PDF invalido ou corrompido.",
            )

        if result.returncode == 8:
            raise ApiError(
                status_code=400,
                code="password_protected_pdf",
                message="PDF protegido por palavra-passe. Remova a protecao primeiro.",
            )

        if result.returncode != 0 or not output_path.exists():
            raise ApiError(
                status_code=500,
                code="ocr_failed",
                message="Falha no processamento OCR.",
                details={"stderr": _trim_process_output(result.stderr)},
            )

        return output_path.read_bytes()


def convert_pdf_to_pdfa(content: bytes, conformance: str) -> bytes:
    """Convert a PDF into the requested PDF/A conformance."""
    pdfa_level = CONFORMANCE_MAP.get(conformance)
    if pdfa_level is None:
        raise ApiError(
            status_code=400,
            code="invalid_conformance",
            message=(
                f"Nível de conformidade inválido: {conformance}. "
                f"Suportados: {', '.join(CONFORMANCE_MAP.keys())}"
            ),
        )

    pdfa_definition_template = next(Path("/usr/share/ghostscript").rglob("PDFA_def.ps"), None)
    icc_profile_path = Path("/usr/share/color/icc/ghostscript/default_rgb.icc")
    if pdfa_definition_template is None or not icc_profile_path.exists():
        raise ApiError(
            status_code=503,
            code="tool_unavailable",
            message="Os recursos Ghostscript para PDF/A não estão disponíveis neste ambiente.",
        )

    template_text = pdfa_definition_template.read_text(encoding="latin-1")

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir_path = Path(tmpdir)
        input_path = tmpdir_path / "input.pdf"
        output_path = tmpdir_path / "output.pdf"
        pdfa_definition_path = tmpdir_path / "PDFA_def.ps"

        input_path.write_bytes(content)
        pdfa_definition_path.write_text(
            template_text.replace("/ICCProfile (srgb.icc)", f"/ICCProfile ({icc_profile_path})"),
            encoding="latin-1",
        )

        result = _run_command(
            [
                "gs",
                f"-dPDFA={pdfa_level}",
                "-dBATCH",
                "-dNOPAUSE",
                "-sDEVICE=pdfwrite",
                "-sColorConversionStrategy=RGB",
                "-dPDFACompatibilityPolicy=1",
                f"--permit-file-read={icc_profile_path}:{pdfa_definition_path}:{input_path}",
                f"-sOutputFile={output_path}",
                str(pdfa_definition_path),
                str(input_path),
            ],
            timeout=55,
            missing_message="Ghostscript não está disponível neste ambiente.",
        )

        if result.returncode != 0 or not output_path.exists():
            stderr = _trim_process_output(result.stderr)
            status_code = 400 if "Password" in stderr or "password" in stderr else 500
            error_code = (
                "password_protected_pdf" if status_code == 400 else "pdfa_conversion_failed"
            )
            message = (
                "PDF protegido por palavra-passe. Remova a protecao primeiro."
                if status_code == 400
                else "PDF/A conversion failed."
            )
            raise ApiError(
                status_code=status_code,
                code=error_code,
                message=message,
                details={"stderr": stderr},
            )

        return output_path.read_bytes()


def _open_pikepdf(content: bytes):
    import pikepdf

    try:
        return pikepdf.open(io.BytesIO(content))
    except pikepdf.PasswordError as exc:
        raise ApiError(
            status_code=400,
            code="password_protected_pdf",
            message="Este PDF ja esta protegido com palavra-passe.",
        ) from exc
    except Exception as exc:
        raise ApiError(
            status_code=400,
            code="invalid_pdf",
            message="Nao foi possivel abrir o PDF. Verifique se o ficheiro e valido.",
        ) from exc


def protect_pdf(content: bytes, password: str) -> bytes:
    """Encrypt a PDF with AES-256 permissions."""
    import pikepdf

    if not password or not password.strip():
        raise ApiError(
            status_code=400,
            code="invalid_password",
            message="A palavra-passe é obrigatória.",
        )

    pdf = _open_pikepdf(content)
    permissions = pikepdf.Permissions(
        accessibility=True,
        extract=False,
        modify_annotation=False,
        modify_assembly=False,
        modify_form=False,
        modify_other=False,
        print_lowres=True,
        print_highres=True,
    )

    try:
        buf = io.BytesIO()
        pdf.save(
            buf,
            encryption=pikepdf.Encryption(
                owner=password,
                user=password,
                R=6,
                aes=True,
                allow=permissions,
            ),
        )
        return buf.getvalue()
    except Exception as exc:
        raise ApiError(
            status_code=500,
            code="protection_failed",
            message="Não foi possível proteger o PDF.",
        ) from exc
    finally:
        pdf.close()


def _set_field_value(field: Any, value: Any) -> None:
    import pikepdf
    from pikepdf.form import (
        CheckboxField,
        ChoiceField,
        MultipleFieldProxy,
        RadioButtonGroup,
        TextField,
    )

    if isinstance(field, MultipleFieldProxy):
        for sub_field in field:
            _set_field_value(sub_field, value)
        return

    if isinstance(field, CheckboxField):
        field.checked = bool(value)
    elif isinstance(field, RadioButtonGroup):
        str_value = str(value)
        for opt in field.options:
            if str(opt.on_value) == str_value or str(opt.on_value) == f"/{str_value}":
                opt.select()
                return
        field.value = pikepdf.Name(f"/{str_value}")
    elif isinstance(field, (ChoiceField, TextField)):
        field.value = str(value)
    else:
        field.value = str(value)


def fill_form_pdf(
    content: bytes,
    field_values: dict[str, Any],
    *,
    strict_unknown_fields: bool,
) -> bytes:
    """Fill an AcroForm PDF and flatten the result."""
    from pikepdf.form import ExtendedAppearanceStreamGenerator, Form

    if not isinstance(field_values, dict) or not field_values:
        raise ApiError(
            status_code=400,
            code="missing_form_values",
            message="Nenhum valor de campo fornecido.",
        )

    pdf = _open_pikepdf(content)
    if pdf.Root.get("/AcroForm") is None:
        pdf.close()
        raise ApiError(
            status_code=400,
            code="missing_form_fields",
            message="Este PDF não contém campos de formulário.",
        )

    try:
        form = Form(pdf, generate_appearances=ExtendedAppearanceStreamGenerator)
        unknown_fields: list[str] = []

        for field_name, value in field_values.items():
            try:
                field = form[field_name]
            except KeyError:
                unknown_fields.append(field_name)
                continue
            _set_field_value(field, value)

        if strict_unknown_fields and unknown_fields:
            raise ApiError(
                status_code=422,
                code="unknown_form_fields",
                message="Alguns nomes de campo não existem no formulário PDF carregado.",
                details={"unknownFields": sorted(unknown_fields)},
            )

        pdf.flatten_annotations("all")
        buf = io.BytesIO()
        pdf.save(buf)
        return buf.getvalue()
    except ApiError:
        raise
    except Exception as exc:
        raise ApiError(
            status_code=500,
            code="form_processing_failed",
            message="Não foi possível processar o formulário.",
        ) from exc
    finally:
        pdf.close()


def redact_pdf(
    content: bytes,
    *,
    strategy: str,
    custom_text: str = "",
    regex_pattern: str = "",
    confirmed_ids: list[str] | None = None,
) -> bytes:
    """Apply PII redaction. If confirmed_ids is None, redact every match.
    If it is a list, redact only matches whose id is in the list (unknown ids
    are silently skipped — the user's preview saw a set; we trust intent)."""
    import pymupdf

    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ApiError(400, "invalid_pdf",
            "Não foi possível abrir o PDF. Verifique se o ficheiro é válido.") from exc

    try:
        # _extract_matches raises 400 password_protected_pdf when applicable
        all_matches = _extract_matches(
            doc, strategy=strategy, custom_text=custom_text, regex_pattern=regex_pattern,
        )

        if confirmed_ids is not None:
            confirmed_set = set(confirmed_ids)
            matches_to_apply = [m for m in all_matches if m.id in confirmed_set]
        else:
            matches_to_apply = all_matches

        if matches_to_apply:
            with REDACTION_LOCK:
                pymupdf.TOOLS.set_small_glyph_heights(True)
                try:
                    pages_with_redactions: set[int] = set()
                    for m in matches_to_apply:
                        page = doc[m.page]
                        page.add_redact_annot(pymupdf.Rect(*m.bbox), fill=(0, 0, 0))
                        pages_with_redactions.add(m.page)

                    for page_idx in pages_with_redactions:
                        doc[page_idx].apply_redactions(
                            images=pymupdf.PDF_REDACT_IMAGE_NONE,
                            graphics=pymupdf.PDF_REDACT_LINE_ART_NONE,
                            text=pymupdf.PDF_REDACT_TEXT_REMOVE,  # explicit: delete characters
                        )
                finally:
                    pymupdf.TOOLS.set_small_glyph_heights(False)

        # CRITICAL: strip residual sensitive data from outline/metadata/hidden text.
        # Without this, the "redacted" PDF can still leak the redacted content via
        # bookmark titles (EU AstraZeneca 2021), metadata (multiple), or OCR
        # invisible layers (Epstein documents 2024). See spec for full citations.
        # Runs UNCONDITIONALLY — even a no-match round-trip must strip metadata.
        doc.scrub(
            attached_files=True,
            embedded_files=True,
            hidden_text=True,
            javascript=True,
            metadata=True,
            xml_metadata=True,
            remove_links=True,
            reset_fields=True,
            reset_responses=True,
            thumbnails=True,
            clean_pages=True,
            redactions=False,   # already applied above with explicit options
            redact_images=0,
        )
        # scrub() does NOT touch the document outline (verified against pymupdf
        # 1.27 docs). Clear the TOC explicitly so bookmark titles cannot leak
        # — this is the exact EU AstraZeneca 2021 failure mode.
        doc.set_toc([])

        return doc.tobytes(garbage=4, deflate=True, clean=True)
    except ApiError:
        raise
    except Exception as exc:
        raise ApiError(500, "redaction_failed",
            "Não foi possível processar a redacção deste PDF.") from exc
    finally:
        doc.close()


def _docx_is_effectively_empty(path: Path) -> bool:
    """True when the produced docx has no non-whitespace text and no tables.

    pdf2docx swallows per-page errors (ignore_page_error=True) and exits 0,
    so a scanned/degraded PDF yields a valid-but-empty docx no exception flags.
    """
    from docx import Document

    document = Document(str(path))
    has_text = any(p.text.strip() for p in document.paragraphs)
    return not has_text and len(document.tables) == 0


def pdf_to_docx(content: bytes) -> bytes:
    """Convert a text-based PDF to an editable .docx via the pdf2docx CLI.

    Guard order is load-bearing: encrypted before page access, page cap before
    the get_text loop, then the scanned/empty-text gate.
    """
    import pymupdf

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.pdf"
        output_path = Path(tmpdir) / "output.docx"
        input_path.write_bytes(content)

        try:
            doc = pymupdf.open(stream=content, filetype="pdf")
        except Exception as exc:
            raise ApiError(
                status_code=400,
                code="invalid_pdf",
                message="Não foi possível abrir o PDF. Verifique se o ficheiro é válido.",
            ) from exc

        try:
            if doc.needs_pass:
                raise ApiError(
                    status_code=400,
                    code="password_protected_pdf",
                    message="Este PDF está protegido por palavra-passe. Desbloqueie-o antes de converter.",
                )
            if doc.page_count > MAX_PAGES:
                raise ApiError(
                    status_code=400,
                    code="too_many_pages",
                    message="PDF demasiado grande. Divida-o antes de converter.",
                )
            total_chars = sum(len("".join(page.get_text().split())) for page in doc)
            if total_chars < doc.page_count * 10:
                raise ApiError(
                    status_code=422,
                    code="scanned_pdf",
                    message=(
                        "Este PDF parece digitalizado (sem texto selecionável). "
                        "Use a ferramenta OCR primeiro e depois converta."
                    ),
                )
        finally:
            doc.close()

        result = _run_command(
            ["pdf2docx", "convert", str(input_path), str(output_path)],
            timeout=55,
            missing_message="Conversor indisponível neste ambiente.",
        )
        if result.returncode != 0 or not output_path.exists():
            logger.error("pdf2docx failed: %s", _trim_process_output(result.stderr))
            raise ApiError(
                status_code=500,
                code="conversion_failed",
                message="Falha na conversão do PDF para Word.",
            )

        if _docx_is_effectively_empty(output_path):
            raise ApiError(
                status_code=422,
                code="scanned_pdf",
                message=(
                    "Não foi possível extrair texto deste PDF. "
                    "Se for digitalizado, use a ferramenta OCR primeiro."
                ),
            )

        return output_path.read_bytes()


def _sheet_title(pno: int, ti: int) -> str:
    """Build a worksheet title for the table `ti` on page `pno`.

    openpyxl forbids ``* ? : / [ ] \\`` and caps titles at 31 chars; this format
    uses none of those and stays well within 31 chars for MAX_PAGES/MAX_TABLES.
    Duplicates are auto-deduped by openpyxl; the [:31] slice is defensive only.
    """
    return f"Pag {pno} Tabela {ti}"[:31]


def pdf_to_xlsx(content: bytes) -> bytes:
    """Extract tables from a text PDF into an .xlsx (one sheet per table).

    Guard order is load-bearing: invalid -> encrypted -> page-cap (all cheap,
    before the loop) -> per-page complexity -> per-table cells-cap. The
    scanned-vs-no-tables decision is deferred until after extraction so a sparse
    legit table is not pre-rejected by the text-density gate. The whole body is
    wrapped in try/except -> ApiError(500) because main.py has no catch-all
    handler, so a raw find_tables()/openpyxl exception would escape as a
    code-less plain-text 500 that the FE cannot render.
    """
    from io import BytesIO

    import pymupdf
    from openpyxl import Workbook

    try:
        doc = pymupdf.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ApiError(
            400,
            "invalid_pdf",
            "Não foi possível abrir o PDF. Verifique se o ficheiro é válido.",
        ) from exc

    try:
        # --- Cheap pre-flight (reuse pdf_to_docx codes/messages) ---
        if doc.needs_pass:
            raise ApiError(
                400,
                "password_protected_pdf",
                "Este PDF está protegido por palavra-passe. Desbloqueie-o antes de converter.",
            )
        if doc.page_count > MAX_PAGES:
            raise ApiError(
                400,
                "too_many_pages",
                "PDF demasiado grande. Divida-o antes de converter.",
            )

        wb = Workbook()
        wb.remove(wb.active)  # start with zero sheets
        n_tables, n_cells = 0, 0
        for pno, page in enumerate(doc, start=1):
            # Complexity: reject a vector-graphics bomb before find_tables' O(n²) clustering.
            if len(page.get_cdrawings()) > MAX_PATHS_PER_PAGE:
                raise ApiError(
                    422,
                    "pdf_too_complex",
                    "Página demasiado complexa para extrair tabelas com segurança.",
                )
            for ti, tab in enumerate(page.find_tables().tables, start=1):
                rows = tab.extract()
                if not rows:
                    continue
                n_cells += sum(len(r) for r in rows)
                if n_cells > MAX_CELLS:  # memory cap
                    raise ApiError(
                        422,
                        "pdf_too_complex",
                        "Demasiadas células para um só ficheiro.",
                    )
                n_tables += 1
                ws = wb.create_sheet(title=_sheet_title(pno, ti))
                for r_idx, row in enumerate(rows, start=1):
                    for c_idx, val in enumerate(row, start=1):
                        cell = ws.cell(
                            row=r_idx,
                            column=c_idx,
                            value=("" if val is None else str(val)),
                        )
                        # Neutralize formula ('f') / error ('e') injection: a cell
                        # like "=HYPERLINK(...)" would otherwise ship as a live
                        # formula, and "=1+1"/"#REF!" would render as an error.
                        if cell.data_type in ("f", "e"):
                            cell.data_type = "s"
                if n_tables >= MAX_TABLES:
                    break
            if n_tables >= MAX_TABLES:
                break

        if n_tables == 0:
            # Decide scanned-vs-no-tables AFTER extraction: a sparse legit table
            # must not be pre-rejected by the text-density gate, and a scanned
            # PDF yields zero tables and lands here anyway.
            total_chars = sum(len("".join(p.get_text().split())) for p in doc)
            if total_chars < doc.page_count * 10:
                raise ApiError(
                    422,
                    "scanned_pdf",
                    "Não foi possível extrair texto deste PDF. "
                    "Se for digitalizado, use a ferramenta OCR primeiro.",
                )
            raise ApiError(
                422,
                "no_tables_detected",
                "Não encontrámos tabelas neste PDF. Esta ferramenta extrai tabelas; "
                "para converter texto corrido use o PDF para Word.",
            )

        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()
    except ApiError:
        raise
    except Exception as exc:
        logger.error("pdf_to_xlsx failed: %s", exc)
        raise ApiError(
            500,
            "conversion_failed",
            "Falha na conversão do PDF para Excel.",
        ) from exc
    finally:
        doc.close()


def build_health_payload() -> dict[str, Any]:
    """Return service health data and dependency versions."""
    try:
        import pymupdf

        pymupdf_version = pymupdf.VersionBind
    except ImportError:
        pymupdf_version = "unavailable"

    try:
        import pikepdf

        pikepdf_version = pikepdf.__version__
    except ImportError:
        pikepdf_version = "unavailable"

    def read_version(command: list[str], timeout: int, *, first_line: bool = False) -> str:
        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=timeout,
            )
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return "unavailable"

        if result.returncode != 0:
            return "unavailable"

        output = result.stdout.strip()
        if first_line:
            output = output.splitlines()[0] if output else ""
        return output or "unavailable"

    return {
        "status": "ok",
        "versions": {
            "pymupdf": pymupdf_version,
            "pikepdf": pikepdf_version,
            "ghostscript": read_version(["gs", "--version"], 10),
            "tesseract": read_version(["tesseract", "--version"], 10, first_line=True),
            "libreoffice": read_version(["soffice", "--headless", "--version"], 30),
        },
    }
